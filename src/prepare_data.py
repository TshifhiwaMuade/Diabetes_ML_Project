from pathlib import Path

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent

# Prefer the recommended project structure first: data/raw and data/processed.
RAW_DIR = BASE_DIR / "data" / "raw"
PROCESSED_DIR = BASE_DIR / "data" / "processed"

# Fallbacks for older/local layouts.
if not RAW_DIR.exists() and (BASE_DIR / "raw").exists():
    RAW_DIR = BASE_DIR / "raw"

if not PROCESSED_DIR.exists() and (BASE_DIR / "processed").exists():
    PROCESSED_DIR = BASE_DIR / "processed"

REPORTS_DIR = BASE_DIR / "reports"
FIGURES_DIR = REPORTS_DIR / "figures"

RAW_FILE = RAW_DIR / "Diabetes_and_LifeStyle_Dataset_.csv"
if not RAW_FILE.exists() and (BASE_DIR / "Diabetes_and_LifeStyle_Dataset_.csv").exists():
    RAW_FILE = BASE_DIR / "Diabetes_and_LifeStyle_Dataset_.csv"

if not RAW_FILE.exists():
    raise FileNotFoundError(
        f"Dataset not found. Expected one of: "
        f"{BASE_DIR / 'data' / 'raw' / 'Diabetes_and_LifeStyle_Dataset_.csv'} or "
        f"{BASE_DIR / 'raw' / 'Diabetes_and_LifeStyle_Dataset_.csv'}"
    )

PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
FIGURES_DIR.mkdir(parents=True, exist_ok=True)

# LOAD DATA
df = pd.read_csv(RAW_FILE)

print("Resolved dataset path:", RAW_FILE)
print("Shape:", df.shape)
print("\nColumns:")
print(df.columns.tolist())

print("\nData types:")
print(df.dtypes)

print("\nFirst 5 rows:")
print(df.head())

print("\nTarget distribution:")
print(df["diabetes_stage"].value_counts())

# COLUMN TYPES
categorical_cols = df.select_dtypes(include=["object"]).columns.tolist()
numeric_cols = df.select_dtypes(exclude=["object"]).columns.tolist()

print("\nCategorical columns:")
print(categorical_cols)

print("\nNumeric columns:")
print(numeric_cols)

print("\nNumber of categorical columns:", len(categorical_cols))
print("Number of numeric columns:", len(numeric_cols))

# FEATURE GROUPS
demographic_features = [
    "Age", "gender", "ethnicity", "education_level",
    "income_level", "employment_status"
]

lifestyle_features = [
    "smoking_status", "alcohol_consumption_per_week",
    "physical_activity_minutes_per_week", "diet_score",
    "sleep_hours_per_day", "screen_time_hours_per_day"
]

medical_history_features = [
    "family_history_diabetes", "hypertension_history",
    "cardiovascular_history"
]

clinical_features = [
    "bmi", "waist_to_hip_ratio", "systolic_bp", "diastolic_bp",
    "heart_rate", "cholesterol_total", "hdl_cholesterol",
    "ldl_cholesterol", "triglycerides", "glucose_fasting",
    "glucose_postprandial", "insulin_level", "hba1c",
    "diabetes_risk_score", "diagnosed_diabetes"
]

target_feature = "diabetes_stage"

expected_columns = (
    demographic_features
    + lifestyle_features
    + medical_history_features
    + clinical_features
    + [target_feature]
)
missing_expected_columns = [col for col in expected_columns if col not in df.columns]
if missing_expected_columns:
    raise KeyError(
        "The following expected columns were not found in the dataset: "
        + ", ".join(missing_expected_columns)
    )

print("\nDemographic features:", demographic_features)
print("\nLifestyle features:", lifestyle_features)
print("\nMedical history features:", medical_history_features)
print("\nClinical features:", clinical_features)
print("\nTarget feature:", target_feature)

# MISSING VALUES AND DUPLICATES
missing_values = df.isnull().sum()
total_missing = missing_values.sum()
duplicate_count = df.duplicated().sum()

print("\nMissing values per column:")
print(missing_values)

print("\nTotal missing values in dataset:")
print(total_missing)

print("\nNumber of duplicate rows:")
print(duplicate_count)

# DESCRIPTIVE STATISTICS
desc_stats = df[numeric_cols].describe().T
target_counts = df["diabetes_stage"].value_counts()
target_percentages = (df["diabetes_stage"].value_counts(normalize=True) * 100).round(2)

print("\nDescriptive statistics for numeric columns:")
print(desc_stats)

print("\nTarget percentage distribution:")
print(target_percentages)

key_features = ["glucose_fasting", "hba1c", "bmi", "triglycerides"]

key_feature_stats = df[key_features].describe(
    percentiles=[0.01, 0.05, 0.25, 0.5, 0.75, 0.95, 0.99]
).T

print("\nDetailed statistics for key features:")
print(key_feature_stats)


# BOXPLOTS
for col in key_features:
    plt.figure(figsize=(8, 4))
    plt.boxplot(df[col], vert=False)
    plt.title(f"Boxplot of {col}")
    plt.xlabel(col)
    plt.tight_layout()
    plt.savefig(FIGURES_DIR / f"{col}_boxplot.png")
    plt.close()

# OUTLIER ANALYSIS
outlier_summary = []

for col in key_features:
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    outlier_count = ((df[col] < lower_bound) | (df[col] > upper_bound)).sum()
    outlier_percent = (outlier_count / len(df)) * 100

    outlier_summary.append({
        "feature": col,
        "Q1": round(Q1, 2),
        "Q3": round(Q3, 2),
        "IQR": round(IQR, 2),
        "lower_bound": round(lower_bound, 2),
        "upper_bound": round(upper_bound, 2),
        "outlier_count": int(outlier_count),
        "outlier_percent": round(outlier_percent, 2)
    })

outlier_df = pd.DataFrame(outlier_summary)

print("\nIQR Outlier Summary:")
print(outlier_df)


# TARGET LEAKAGE CHECKS
diagnosed_crosstab = pd.crosstab(df["diagnosed_diabetes"], df["diabetes_stage"])
diagnosed_crosstab_pct = pd.crosstab(
    df["diagnosed_diabetes"],
    df["diabetes_stage"],
    normalize="index"
) * 100

risk_score_by_stage = df.groupby("diabetes_stage")["diabetes_risk_score"].describe()

print("\nCrosstab: diagnosed_diabetes vs diabetes_stage")
print(diagnosed_crosstab)

print("\nRow percentages for diagnosed_diabetes vs diabetes_stage")
print(diagnosed_crosstab_pct)

print("\nAverage diabetes_risk_score by diabetes_stage")
print(risk_score_by_stage)


# CLEANED DATASETS
df_clean_full = df.copy()

cols_to_drop = ["diagnosed_diabetes", "diabetes_risk_score"]
df_clean_model = df.drop(columns=cols_to_drop).copy()

print("\nFull cleaned dataset shape:")
print(df_clean_full.shape)

print("\nNo-leakage modelling dataset shape:")
print(df_clean_model.shape)

print("\nColumns removed from modelling dataset:")
print(cols_to_drop)

print("\nRemaining columns in modelling dataset:")
print(df_clean_model.columns.tolist())

clean_full_path = PROCESSED_DIR / "cleaned_full_dataset.csv"
clean_model_path = PROCESSED_DIR / "cleaned_model_dataset.csv"

df_clean_full.to_csv(clean_full_path, index=False)
df_clean_model.to_csv(clean_model_path, index=False)

print("\nDatasets saved successfully.")
print("Saved full cleaned dataset to:", clean_full_path)
print("Saved modelling dataset to:", clean_model_path)


# WORD DOCUMENT HELPERS
def add_dataframe_table(doc, dataframe, title=None, max_rows=None):
    if title:
        doc.add_paragraph(title)

    df_to_use = dataframe.copy()
    if max_rows is not None:
        df_to_use = df_to_use.head(max_rows)

    table = doc.add_table(rows=1, cols=len(df_to_use.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df_to_use.columns):
        hdr_cells[i].text = str(col_name)

    for _, row in df_to_use.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.add_paragraph("")


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


# CREATE WORD DOCUMENT
doc = Document()

doc.add_heading("Nicholas – Data Understanding & Preparation", level=1)
doc.add_paragraph(
    "Diabetes Risk Segmentation & Decision Support System\n"
    "BC Analytics | MLG382 Guided Project 2026"
)

# Dataset overview
doc.add_heading("1. Dataset Overview", level=2)
doc.add_paragraph(f"The dataset contains {df.shape[0]:,} rows and {df.shape[1]} columns.")
doc.add_paragraph(f"There are {len(numeric_cols)} numeric columns and {len(categorical_cols)} categorical columns.")
doc.add_paragraph(f"The target feature is '{target_feature}'.")

overview_df = pd.DataFrame({
    "Metric": ["Rows", "Columns", "Numeric Columns", "Categorical Columns", "Target Feature"],
    "Value": [df.shape[0], df.shape[1], len(numeric_cols), len(categorical_cols), target_feature]
})
add_dataframe_table(doc, overview_df)

# Feature groups
doc.add_heading("2. Feature Grouping", level=2)
doc.add_paragraph("The dataset was grouped into the following feature categories:")

doc.add_paragraph("Demographic Features:")
add_bullet_list(doc, demographic_features)

doc.add_paragraph("Lifestyle Features:")
add_bullet_list(doc, lifestyle_features)

doc.add_paragraph("Medical History Features:")
add_bullet_list(doc, medical_history_features)

doc.add_paragraph("Clinical Features:")
add_bullet_list(doc, clinical_features)

# Data quality checks
doc.add_heading("3. Data Quality Checks", level=2)
doc.add_paragraph(f"Total missing values in the dataset: {int(total_missing)}")
doc.add_paragraph(f"Duplicate rows in the dataset: {int(duplicate_count)}")

quality_df = pd.DataFrame({
    "Check": ["Total Missing Values", "Duplicate Rows"],
    "Result": [int(total_missing), int(duplicate_count)]
})
add_dataframe_table(doc, quality_df)

doc.add_paragraph(
    "Interpretation: The dataset contains no missing values and no duplicate rows. "
    "This means no imputation or duplicate removal was required at this stage."
)

# Target distribution
doc.add_heading("4. Target Distribution", level=2)

target_dist_df = pd.DataFrame({
    "diabetes_stage": target_counts.index,
    "count": target_counts.values,
    "percentage": target_percentages.values
})
add_dataframe_table(doc, target_dist_df)

doc.add_paragraph(
    "Interpretation: The target variable is highly imbalanced. Type 2 and Pre-Diabetes "
    "make up most of the dataset, while Type 1 and Gestational classes are very small."
)

# Descriptive statistics
doc.add_heading("5. Descriptive Statistics", level=2)
doc.add_paragraph("Summary statistics for key numeric features are shown below.")
add_dataframe_table(
    doc,
    key_feature_stats.reset_index().rename(columns={"index": "feature"}),
    max_rows=None
)

# Outlier analysis
doc.add_heading("6. Outlier Analysis", level=2)
doc.add_paragraph(
    "Boxplots and the IQR method were used to inspect outliers in glucose_fasting, "
    "hba1c, bmi, and triglycerides."
)

add_dataframe_table(doc, outlier_df)

doc.add_paragraph(
    "Interpretation: Small proportions of outliers were detected in the four key clinical "
    "features. Because these variables are clinically relevant to diabetes severity, the "
    "outliers were treated as potentially meaningful rather than automatically removed."
)

for col in key_features:
    image_path = FIGURES_DIR / f"{col}_boxplot.png"
    if image_path.exists():
        doc.add_heading(f"Boxplot: {col}", level=3)
        doc.add_picture(str(image_path), width=Inches(6))

# Leakage checks
doc.add_heading("7. Leakage Checks and Handoff Notes", level=2)
doc.add_paragraph("The following checks were performed to identify possible target leakage.")

doc.add_paragraph("Crosstab: diagnosed_diabetes vs diabetes_stage")
add_dataframe_table(doc, diagnosed_crosstab.reset_index())

doc.add_paragraph("Row percentages: diagnosed_diabetes vs diabetes_stage")
add_dataframe_table(doc, diagnosed_crosstab_pct.reset_index().round(2))

doc.add_paragraph("diabetes_risk_score by diabetes_stage")
add_dataframe_table(doc, risk_score_by_stage.reset_index().round(2))

doc.add_paragraph(
    "Interpretation: The feature 'diagnosed_diabetes' appears to leak the target because it is "
    "strongly aligned with diabetes_stage. The feature 'diabetes_risk_score' is also strongly "
    "target-related and should be reviewed carefully before modelling."
)

# Cleaned dataset handoff
doc.add_heading("8. Cleaned Dataset Handoff", level=2)
doc.add_paragraph(f"Full cleaned dataset shape: {df_clean_full.shape}")
doc.add_paragraph(f"No-leakage modelling dataset shape: {df_clean_model.shape}")
doc.add_paragraph(f"Columns removed from modelling dataset: {', '.join(cols_to_drop)}")

handoff_df = pd.DataFrame({
    "Dataset": ["cleaned_full_dataset.csv", "cleaned_model_dataset.csv"],
    "Rows": [df_clean_full.shape[0], df_clean_model.shape[0]],
    "Columns": [df_clean_full.shape[1], df_clean_model.shape[1]]
})
add_dataframe_table(doc, handoff_df)

doc.add_paragraph(
    "Interpretation: Two datasets were saved. The full cleaned dataset preserves all columns for "
    "reference, while the modelling dataset removes the most suspicious leakage-related variables."
)

# Final summary
doc.add_heading("9. Summary", level=2)
doc.add_paragraph(
    "Data understanding showed that the dataset is structurally clean, with no missing values "
    "or duplicate rows. The target is highly imbalanced, and the main clinical variables contain "
    "small proportions of plausible outliers. Leakage analysis identified diagnosed_diabetes as a "
    "strong leakage feature, while diabetes_risk_score was also flagged for caution. Cleaned "
    "datasets were saved for downstream handoff."
)

# Save Word document
word_output_path = REPORTS_DIR / "Nicholas_Data_Understanding_Preparation_Evidence.docx"
doc.save(word_output_path)

print(f"\nWord document saved successfully: {word_output_path}")
